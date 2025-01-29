import sys
import time
import json
from pathlib import Path
import requests
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QPushButton, QLabel,
    QProgressBar, QTextEdit
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal

# Hadith books and sections
hadith_books = {
    "bukhari": 97, "muslim": 56, "nasai": 51, "abudawud": 43,
    "tirmidhi": 49, "ibnmajah": 37, "malik": 61, "ahmad": 71,
    "adab": 57, "shamail": 56, "bulugh": 16
}

BASE_URL = "https://sunnah.com/"
OUTPUT_DIR = Path("hadith_documents")
OUTPUT_DIR.mkdir(exist_ok=True)  # Ensure the folder exists
PROGRESS_FILE = "progress.json"  # File to save progress


class HadithScraper(QThread):
    """Background thread to scrape Hadiths."""
    log_signal = pyqtSignal(str)
    status_signal = pyqtSignal(str)
    progress_signal = pyqtSignal(int, int)  # (overall, book)
    finished_signal = pyqtSignal()

    def __init__(self):
        super().__init__()
        self.stop_scraping = False
        self.current_book = None
        self.current_section = 1
        self.load_progress()

    def run(self):
        """Main scraping function."""
        self.stop_scraping = False
        driver = self.init_driver()

        total_books = len(hadith_books)
        book_keys = list(hadith_books.keys())

        for book_index, book in enumerate(book_keys):
            if book_index < self.current_book:
                continue  # Skip already scraped books

            sections = hadith_books[book]
            self.log_signal.emit(f"\n📖 Scraping {book} ({sections} sections)...")
            self.status_signal.emit(f"Scraping {book}...")

            for section in range(self.current_section, sections + 1):
                if self.stop_scraping:
                    self.save_progress(book_index, section)
                    self.log_signal.emit("🛑 Scraping stopped by user. Progress saved.")
                    return

                hadith_url = f"{BASE_URL}{book}/{section}"
                self.log_signal.emit(f"🔗 Fetching: {hadith_url}")

                page_source = self.fetch_page(hadith_url, driver)
                if not page_source:
                    self.log_signal.emit(f"❌ Skipping {hadith_url}")
                    continue

                soup = BeautifulSoup(page_source, "html.parser")
                title, combined_hadith = self.extract_hadith(soup)
                self.save_hadith(title, combined_hadith, book, section)

                self.progress_signal.emit(book_index, section)
                time.sleep(2)  # Prevent server overload

            self.current_section = 1  # Reset section when moving to next book

        if driver:
            driver.quit()

        self.clear_progress()  # Clear progress when finished
        self.finished_signal.emit()
        self.status_signal.emit("✅ Scraping Completed")
        self.log_signal.emit("✅ Scraping Completed")

    def init_driver(self):
        """Initialize Selenium WebDriver."""
        options = Options()
        options.add_argument("--headless")
        options.add_argument("--disable-gpu")
        options.add_argument("--log-level=3")
        options.add_argument("--user-agent=Mozilla/5.0")

        try:
            driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
            return driver
        except Exception as e:
            self.log_signal.emit(f"❌ WebDriver Error: {e}")
            return None

    def fetch_page(self, url, driver=None):
        """Fetch page source with Selenium or Requests."""
        try:
            if driver:
                driver.get(url)
                return driver.page_source
            else:
                headers = {'User-Agent': 'Mozilla/5.0'}
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()
                return response.text
        except requests.exceptions.RequestException as e:
            self.log_signal.emit(f"❌ Request failed: {url}")
            return None

    def extract_hadith(self, soup):
        """Extract Hadith content from the page."""
        title_tag = soup.find("title")
        title = title_tag.text.strip() if title_tag else "No Title Found"
        hadith_english = soup.find_all("div", "english_hadith_full")
        hadith_arabic = soup.find_all("div", "arabic_hadith_full")
        hadith_reference = soup.find_all("div", "bottomItems")

        return title, zip(hadith_english, hadith_arabic, hadith_reference)

    def save_hadith(self, title, combined_hadith, book, section):
        """Save Hadith to a Word document."""
        doc_path = OUTPUT_DIR / f"{book}_{section}.docx"
        document = Document()
        document.add_heading(title, level=1)

        for english, arabic, reference in combined_hadith:
            document.add_paragraph(english.text.strip(), style="Normal")
            document.add_paragraph(arabic.text.strip(), style="Normal")
            document.add_paragraph(reference.text.strip(), style="Normal")
            document.add_paragraph("-" * 40)

        document.save(doc_path)
        self.log_signal.emit(f"✔ Saved: {doc_path}")

    def load_progress(self):
        """Load progress from file."""
        if Path(PROGRESS_FILE).exists():
            with open(PROGRESS_FILE, "r") as file:
                progress = json.load(file)
                self.current_book = progress.get("current_book", 0)
                self.current_section = progress.get("current_section", 1)
        else:
            self.current_book = 0
            self.current_section = 1

    def save_progress(self, book_index, section):
        """Save current progress to file."""
        progress = {"current_book": book_index, "current_section": section}
        with open(PROGRESS_FILE, "w") as file:
            json.dump(progress, file)

    def clear_progress(self):
        """Remove progress file when scraping completes."""
        if Path(PROGRESS_FILE).exists():
            Path(PROGRESS_FILE).unlink()

    def stop(self):
        """Stop scraping process and save progress."""
        self.stop_scraping = True


class HadithScraperGUI(QWidget):
    """Main GUI Application."""

    def __init__(self):
        super().__init__()
        self.init_ui()
        self.scraper_thread = None

    def init_ui(self):
        """Setup GUI layout."""
        self.setWindowTitle("Hadith Scraper")
        self.setGeometry(100, 100, 500, 400)
        layout = QVBoxLayout()

        self.status_label = QLabel("Status: Waiting to start...")
        layout.addWidget(self.status_label)

        self.overall_progress = QProgressBar()
        self.overall_progress.setMaximum(len(hadith_books))
        layout.addWidget(self.overall_progress)

        self.book_progress = QProgressBar()
        layout.addWidget(self.book_progress)

        self.log_output = QTextEdit()
        self.log_output.setReadOnly(True)
        layout.addWidget(self.log_output)

        self.start_button = QPushButton("Start Scraping")
        self.start_button.clicked.connect(self.start_scraping)
        layout.addWidget(self.start_button)

        self.stop_button = QPushButton("Stop")
        self.stop_button.setEnabled(False)
        self.stop_button.clicked.connect(self.stop_scraping)
        layout.addWidget(self.stop_button)

        self.setLayout(layout)

    def start_scraping(self):
        """Start the scraper in a new thread."""
        self.scraper_thread = HadithScraper()
        self.scraper_thread.log_signal.connect(self.log_output.append)
        self.scraper_thread.status_signal.connect(self.status_label.setText)
        self.scraper_thread.progress_signal.connect(self.update_progress)
        self.scraper_thread.finished_signal.connect(self.scraping_finished)
        self.scraper_thread.start()

        self.start_button.setEnabled(False)
        self.stop_button.setEnabled(True)

    def stop_scraping(self):
        """Stop the scraper."""
        self.scraper_thread.stop()
        self.stop_button.setEnabled(False)

    def update_progress(self, overall, book):
        self.overall_progress.setValue(overall)
        self.book_progress.setValue(book)

    def scraping_finished(self):
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)


app = QApplication(sys.argv)
window = HadithScraperGUI()
window.show()
sys.exit(app.exec())