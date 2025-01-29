from bs4 import BeautifulSoup
import requests
from pathlib import Path
from docx import Document
import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager

# Dictionary of hadith books and sections
hadith_books = {
    "bukhari": 97, "muslim": 56, "nasai": 51, "abudawud": 43,
    "tirmidhi": 49, "ibnmajah": 37, "malik": 61, "ahmad": 71,
    "adab": 57, "shamail": 56, "bulugh": 16
}

BASE_URL = 'https://sunnah.com/'
OUTPUT_DIR = Path("hadith_documents")  # Saves output in a dedicated folder
OUTPUT_DIR.mkdir(exist_ok=True)  # Create if not exists

# Set up Selenium WebDriver
def init_driver():
    options = Options()
    options.add_argument("--headless")  # Run in headless mode (no GUI)
    options.add_argument("--disable-gpu")
    options.add_argument("--log-level=3")  # Suppress logs
    options.add_argument("--user-agent=Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36")
    
    try:
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        return driver
    except Exception as e:
        print(f"Error initializing WebDriver: {e}")
        return None

# Function to fetch page source using Selenium or Requests
def fetch_page(url, driver=None):
    try:
        if driver:
            driver.get(url)
            return driver.page_source
        else:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            return response.text
    except requests.exceptions.RequestException as e:
        print(f"Request failed for {url}: {e}")
        return None

# Function to extract Hadith content
def extract_hadith(soup):
    title = soup.find("title").text.strip()
    hadith_english = soup.find_all("div", "english_hadith_full")
    hadith_arabic = soup.find_all("div", "arabic_hadith_full")
    hadith_reference = soup.find_all("div", "bottomItems")
    
    combined_hadith = zip(hadith_english, hadith_arabic, hadith_reference)
    return title, combined_hadith

# Function to save Hadith to a Word document
def save_hadith(title, combined_hadith, book, section):
    doc_path = OUTPUT_DIR / f"{book}_{section}.docx"
    document = Document()
    document.add_heading(title, level=1)

    for english, arabic, reference in combined_hadith:
        document.add_paragraph(english.text.strip(), style="Normal")
        document.add_paragraph(arabic.text.strip(), style="Normal")
        document.add_paragraph(reference.text.strip(), style="Normal")
        document.add_paragraph("-" * 40)

    document.save(doc_path)
    print(f"✔ Saved: {doc_path}")

# Main function to scrape Hadith books
def scrape_hadith_books():
    driver = init_driver()
    
    for book, sections in hadith_books.items():
        print(f"\n📖 Scraping {book} ({sections} sections)")

        for section in range(1, sections + 1):
            hadith_url = f"{BASE_URL}{book}/{section}"
            print(f"🔗 Fetching: {hadith_url}")

            page_source = fetch_page(hadith_url, driver)
            if not page_source:
                print(f"❌ Skipping {hadith_url} due to failed fetch.")
                continue

            soup = BeautifulSoup(page_source, "html.parser")
            title, combined_hadith = extract_hadith(soup)
            save_hadith(title, combined_hadith, book, section)
            
            time.sleep(2)  # Prevent server overload

    if driver:
        driver.quit()

# Run the script
if __name__ == "__main__":
    scrape_hadith_books()