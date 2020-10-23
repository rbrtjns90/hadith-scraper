from bs4 import BeautifulSoup
import requests
from pathlib import Path
from docx import Document
import time
from selenium import webdriver

'''
Hadith with a one page structure

nawawi40, qudsi40, riyadussalihin, mishkat

'''
hadith_books = {"bukhari": 97, "muslim": 56, "nasai": 51, "abudawud": 43,
                "tirmidhi": 49, "ibnmajah": 37, "malik": 61, "ahmad": 71,
                "adab": 57, "shamail": 56, "bulugh": 16,
                }

url = 'https://sunnah.com/'
# Can alternatively use requests or selenium if blocked from scraping activity

driver = webdriver.Chrome() #Comment webdriver elements and uncomment requests to use request (Will get a speed boost)

for book, section in hadith_books.items():
    print("Book and Section")
    print(book, section)
    print("")
    for i in range(1, section):
        built_url = str(url) + str(book) + "/" + str(i)
        print(built_url)
        driver.get(built_url) 
        #headers={'User-Agent':'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36'}
        #r = requests.get(built_url, headers=headers)
        #page_data = r.text
        #print(r.text)
        #soup = BeautifulSoup(r.text, "html.parser")
        
        soup = BeautifulSoup(driver.page_source, "html.parser")
        document = Document()

        title_page = soup.find("title")

        print(title_page.text)

        document.add_heading(title_page.text)
        hadith_english = soup.find_all("div", "english_hadith_full")
        hadith_arabic = soup.find_all("div", "arabic_hadith_full")
        hadith_reference = soup.find_all("div", 'bottomItems')
        combined_hadith = zip(hadith_english, hadith_arabic, hadith_reference)

        for english, arabic, reference in combined_hadith:
            print(english, arabic, reference)
            document.add_paragraph(english.text)
            document.add_paragraph(arabic.text)
            document.add_paragraph(reference.text)
            document.add_paragraph("------------------------")


        document.save(str(book) + str(i) + '.docx')
        time.sleep(3) #To not DOS the server
